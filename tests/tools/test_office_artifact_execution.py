"""Real DOCX/XLSX creation through the public ``exec`` contract."""

from __future__ import annotations

from collections.abc import Callable
from pathlib import Path

import pytest

from deeptutor.services.path_service import PathService
from deeptutor.services.workspace.execution import prepare_workspace_execution_env
from deeptutor.tools.exec_tool import ExecTool
from deeptutor.tools.workspace import WorkspacePresentTool


def _workspace_service(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import deeptutor.services.workspace as workspace_package
    import deeptutor.services.workspace.service as workspace_module
    from deeptutor.services.workspace.service import ContentWorkspaceService
    import deeptutor.tools.workspace as workspace_tools

    path_service = PathService(workspace_root=tmp_path / "runtime-data")
    monkeypatch.setattr(workspace_module, "get_path_service", lambda: path_service)
    service = ContentWorkspaceService()
    binding = service.default_binding()
    binding.root.mkdir(parents=True)
    service.current_binding(ensure_output=True)
    monkeypatch.setattr(workspace_package, "get_content_workspace_service", lambda: service)
    monkeypatch.setattr(workspace_tools, "get_content_workspace_service", lambda: service)
    return service, binding


async def _run_real_python(
    *,
    code: str,
    turn_name: str,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    seed_workspace: Callable[[Path], None] | None = None,
):
    import deeptutor.services.sandbox as sandbox_package
    from deeptutor.services.sandbox.backends import RestrictedSubprocessBackend

    service, binding = _workspace_service(tmp_path, monkeypatch)
    if seed_workspace is not None:
        seed_workspace(binding.root)
    turn_dir = binding.root / "outputs" / "chat" / "office-e2e" / turn_name
    code_workdir = turn_dir / "code_runs"
    backend = RestrictedSubprocessBackend()

    class DirectSandboxService:
        async def run(self, request, *, user_id: str):
            return await backend.exec(request)

    monkeypatch.setattr(sandbox_package, "get_sandbox_service", lambda: DirectSandboxService())
    result = await ExecTool().execute(
        code=code,
        language="python",
        _sandbox_code_workdir=str(code_workdir),
        _sandbox_env=prepare_workspace_execution_env(turn_dir, workspace_root=binding.root),
        _workspace_id=binding.workspace_id,
        _workspace_root=str(binding.root),
    )
    return result, service, binding, turn_dir, code_workdir


@pytest.mark.asyncio
async def test_exec_reads_existing_workspace_binary_through_symbolic_root(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from docx import Document

    def seed(root: Path) -> None:
        source = root / "documents" / "source.docx"
        source.parent.mkdir(parents=True)
        doc = Document()
        doc.add_paragraph("Existing workspace document")
        doc.save(source)

    result, _service, binding, _turn_dir, code_workdir = await _run_real_python(
        code=(
            "import os\n"
            "from pathlib import Path\n"
            "from docx import Document\n"
            "source = Path(os.environ['DEEPTUTOR_WORKSPACE_ROOT']) / "
            "'documents/source.docx'\n"
            "doc = Document(source)\n"
            "assert doc.paragraphs[0].text == 'Existing workspace document'\n"
            "doc.add_paragraph('Edited safely under outputs')\n"
            "doc.save('revised.docx')\n"
            "print('revised.docx')\n"
        ),
        turn_name="existing-docx",
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
        seed_workspace=seed,
    )

    revised = code_workdir / "revised.docx"
    assert result.success is True
    assert revised.is_file()
    assert (binding.root / "documents" / "source.docx").is_file()
    assert str(binding.root) not in result.content
    assert result.metadata["artifacts"][0]["filename"] == "revised.docx"


@pytest.mark.asyncio
async def test_docx_creation_is_collected_and_presentable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    result, service, binding, turn_dir, code_workdir = await _run_real_python(
        code=(
            "from docx import Document\n"
            "import zipfile\n"
            "out = 'deeptutor_note.docx'\n"
            "doc = Document()\n"
            "doc.add_heading('DeepTutor note', level=0)\n"
            "doc.add_paragraph('Created through the unified exec tool.')\n"
            "doc.save(out)\n"
            "check = Document(out)\n"
            "assert check.paragraphs[0].text == 'DeepTutor note'\n"
            "with zipfile.ZipFile(out) as package:\n"
            "    assert package.testzip() is None\n"
            "print(out)\n"
        ),
        turn_name="docx",
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
    )

    artifact = result.metadata["artifacts"][0]
    generated = code_workdir / "deeptutor_note.docx"
    assert result.success is True
    assert generated.stat().st_size > 1_000
    assert generated.is_relative_to(binding.root / "outputs")
    assert artifact["relative_path"].startswith("outputs/")
    assert artifact["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert result.metadata["workspace_items"] == []
    assert result.sources == []
    assert "not yet presented" in result.content

    presented = await WorkspacePresentTool().execute(
        items=[{"path": artifact["relative_path"]}],
        _workspace_id=binding.workspace_id,
    )
    published_path, item = service.resolve_published_item(
        binding.workspace_id,
        presented.metadata["workspace_items"][0]["workspace_item_id"],
    )
    assert presented.success is True
    assert item.relative_path == artifact["relative_path"]
    assert item.url.startswith("/files/workspace-items/")
    assert published_path.read_bytes() == generated.read_bytes()
    assert all(path.is_relative_to(binding.root) for path in turn_dir.rglob("*"))


@pytest.mark.asyncio
async def test_xlsx_creation_is_collected_and_presentable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    result, service, binding, turn_dir, code_workdir = await _run_real_python(
        code=(
            "from openpyxl import Workbook, load_workbook\n"
            "import zipfile\n"
            "out = 'deeptutor_scores.xlsx'\n"
            "wb = Workbook()\n"
            "ws = wb.active\n"
            "ws.title = 'Scores'\n"
            "ws.append(['Learner', 'Score'])\n"
            "ws.append(['Ada', 95])\n"
            "ws.append(['Lin', 88])\n"
            "ws['B4'] = 183\n"
            "wb.save(out)\n"
            "check = load_workbook(out, data_only=False)\n"
            "assert check['Scores']['B2'].value == 95\n"
            "with zipfile.ZipFile(out) as package:\n"
            "    assert package.testzip() is None\n"
            "print(out)\n"
        ),
        turn_name="xlsx",
        tmp_path=tmp_path,
        monkeypatch=monkeypatch,
    )

    artifact = result.metadata["artifacts"][0]
    generated = code_workdir / "deeptutor_scores.xlsx"
    assert result.success is True
    assert generated.stat().st_size > 1_000
    assert generated.is_relative_to(binding.root / "outputs")
    assert artifact["relative_path"].startswith("outputs/")
    assert artifact["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert result.metadata["workspace_items"] == []
    assert result.sources == []
    assert "not yet presented" in result.content

    presented = await WorkspacePresentTool().execute(
        items=[{"path": artifact["relative_path"]}],
        _workspace_id=binding.workspace_id,
    )
    published_path, item = service.resolve_published_item(
        binding.workspace_id,
        presented.metadata["workspace_items"][0]["workspace_item_id"],
    )
    assert presented.success is True
    assert item.relative_path == artifact["relative_path"]
    assert item.url.startswith("/files/workspace-items/")
    assert published_path.read_bytes() == generated.read_bytes()
    assert all(path.is_relative_to(binding.root) for path in turn_dir.rglob("*"))


def test_office_skills_reuse_the_shared_workspace_contract() -> None:
    root = Path(__file__).parents[2] / "deeptutor" / "skills" / "builtin"
    forbidden_tool_names = ("code_execution", "run_code", "code_execute")

    for name in ("docx", "xlsx"):
        text = (root / name / "SKILL.md").read_text(encoding="utf-8")
        assert all(tool_name not in text for tool_name in forbidden_tool_names)
        assert "Use `exec` with complete Python source" in text
        assert "`language: python`" in text
        assert "**User workspace**" in text
        assert "DEEPTUTOR_WORKSPACE_ROOT" not in text
        assert "workspace_present" not in text
        assert "python -c" not in text


def test_docx_xlsx_dependencies_cover_pip_source_cli_and_docker() -> None:
    root = Path(__file__).parents[2]
    manifests = {
        "pip/source": root / "pyproject.toml",
        "pip cli": root / "packaging" / "deeptutor-cli" / "pyproject.toml",
        "requirements": root / "requirements" / "cli.txt",
        "docker runner": root / "Dockerfile.runner",
    }
    for deployment, path in manifests.items():
        text = path.read_text(encoding="utf-8").lower()
        assert "python-docx" in text, deployment
        assert "openpyxl" in text, deployment

    frontend = (root / "web" / "package.json").read_text(encoding="utf-8")
    assert '"docx-preview"' in frontend
    assert '"exceljs"' in frontend
